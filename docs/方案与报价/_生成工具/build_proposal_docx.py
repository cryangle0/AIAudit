# -*- coding: utf-8 -*-
"""
爱对 - 多平台店铺财务自动对账系统 项目方案 docx 生成器

参考自 E:\\angsa\\angsa_data\\AITools\\create_word(排版).py 的排版风格
保持中文字体一致、章节分页、表格居中
"""

import os
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, font_name='宋体', font_size=None, bold=None, color=None):
    """统一字体设置（含中文）"""
    run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)


def remove_heading_border(heading):
    """移除标题下边框（默认 Word heading 自带的横线）"""
    pPr = heading._element.get_or_add_pPr()
    for child in list(pPr):
        if child.tag.endswith('pBdr'):
            pPr.remove(child)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'none')
    bottom.set(qn('w:sz'), '0')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_markdown(md_content):
    """简化 markdown 解析（标题/段落/列表/表格/代码块/水平线）"""
    lines = md_content.split('\n')
    elements = []
    in_code = False
    code_buf = []
    table_buf = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                elements.append(('code', '\n'.join(code_buf)))
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # 表格
        if stripped.startswith('|') and stripped.endswith('|'):
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if cells:
                table_buf.append(cells)
            i += 1
            continue
        elif table_buf:
            elements.append(('table', table_buf.copy()))
            table_buf = []

        # 标题
        if re.match(r'^# [^#]', line):
            elements.append(('h1', line[2:].strip())); i += 1; continue
        if re.match(r'^## [^#]', line):
            elements.append(('h2', line[3:].strip())); i += 1; continue
        if re.match(r'^### [^#]', line):
            elements.append(('h3', line[4:].strip())); i += 1; continue
        if re.match(r'^#### [^#]', line):
            elements.append(('h4', line[5:].strip())); i += 1; continue

        # 列表
        if line.strip().startswith('- '):
            elements.append(('list', line.strip()[2:])); i += 1; continue
        if re.match(r'^\d+\. ', line.strip()):
            m = re.match(r'^\d+\. (.+)', line.strip())
            if m:
                elements.append(('numlist', m.group(1)))
            i += 1
            continue

        # 水平线
        if stripped == '---':
            elements.append(('hr', '')); i += 1; continue

        # 普通段落
        if stripped:
            elements.append(('paragraph', stripped))

        i += 1

    if table_buf:
        elements.append(('table', table_buf))

    return elements


def add_paragraph_with_inline(doc, text, style='paragraph'):
    """添加段落，支持 **加粗** 内联"""
    p = doc.add_paragraph()
    if style == 'paragraph':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.6
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, '宋体', 11, bold=True)
        elif part:
            run = p.add_run(part)
            set_run_font(run, '宋体', 11)


PAGE_BREAK_TITLES = ['第一章', '第二章', '第三章', '第四章', '第五章', '第六章', '附录']


def build_docx(md_file, output_file):
    base_dir = os.path.dirname(os.path.abspath(md_file))

    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()

    # 页面设置（A4 + 标准页边距）
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # 默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    elements = parse_markdown(md_content)

    print(f'解析到 {len(elements)} 个元素')
    types = {}
    for e in elements:
        types[e[0]] = types.get(e[0], 0) + 1
    print(f'类型统计: {types}')

    for elem in elements:
        et = elem[0]

        if et == 'h1':
            h = doc.add_heading(elem[1], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            remove_heading_border(h)
            for run in h.runs:
                set_run_font(run, '黑体', 22, bold=True, color=RGBColor(0, 0, 0))

        elif et == 'h2':
            title = elem[1]
            if any(title.startswith(t) for t in PAGE_BREAK_TITLES):
                doc.add_page_break()
            h = doc.add_heading(title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in h.runs:
                set_run_font(run, '黑体', 16, bold=True, color=RGBColor(0x2C, 0x5F, 0x9E))

        elif et == 'h3':
            h = doc.add_heading(elem[1], level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in h.runs:
                set_run_font(run, '黑体', 14, bold=True, color=RGBColor(0, 0, 0))

        elif et == 'h4':
            h = doc.add_heading(elem[1], level=3)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in h.runs:
                set_run_font(run, '黑体', 12, bold=True, color=RGBColor(0x33, 0x33, 0x33))

        elif et == 'paragraph':
            add_paragraph_with_inline(doc, elem[1])

        elif et == 'list':
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 1.5
            parts = re.split(r'(\*\*[^*]+\*\*)', elem[1])
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, '宋体', 11, bold=True)
                elif part:
                    run = p.add_run(part)
                    set_run_font(run, '宋体', 11)

        elif et == 'numlist':
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.line_spacing = 1.5
            parts = re.split(r'(\*\*[^*]+\*\*)', elem[1])
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, '宋体', 11, bold=True)
                elif part:
                    run = p.add_run(part)
                    set_run_font(run, '宋体', 11)

        elif et == 'code':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            for line in elem[1].split('\n'):
                run = p.add_run(line + '\n')
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

        elif et == 'table':
            rows = elem[1]
            if rows:
                num_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, row_data in enumerate(rows):
                    for j, cell_text in enumerate(row_data):
                        if j < num_cols:
                            cell = table.rows[i].cells[j]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run(cell_text)
                            set_run_font(run, '宋体', 10, bold=(i == 0))
                            if i == 0:
                                shading = OxmlElement('w:shd')
                                shading.set(qn('w:fill'), 'D9E8F5')
                                cell._tc.get_or_add_tcPr().append(shading)
                doc.add_paragraph()

        elif et == 'hr':
            pass

    doc.save(output_file)
    print(f'\n[OK] Word 文档已生成：{output_file}')


if __name__ == '__main__':
    here = os.path.dirname(os.path.abspath(__file__))
    parent = os.path.dirname(here)  # 方案与报价 目录
    md_file = os.path.join(parent, '爱对-项目方案.md')
    out_file = os.path.join(parent, '爱对-项目方案.docx')
    build_docx(md_file, out_file)
